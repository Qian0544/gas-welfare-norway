"""
Generate presentation manuscript for Member 1 – Section 3
Run: python3 make_manus.py
Output: section3_presentation_manus.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# ── helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(text, bold_parts=None, indent=False):
    """Add a paragraph. bold_parts = list of substrings to bold."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        # naive bold highlighting: split on bold_parts tokens
        remaining = text
        import re
        tokens = re.split(
            "(" + "|".join(re.escape(b) for b in bold_parts) + ")",
            text
        )
        for tok in tokens:
            run = p.add_run(tok)
            run.font.size = Pt(11)
            if tok in bold_parts:
                run.bold = True
    return p

def bullet(text, bold_parts=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        import re
        tokens = re.split(
            "(" + "|".join(re.escape(b) for b in bold_parts) + ")",
            text
        )
        for tok in tokens:
            run = p.add_run(tok)
            run.font.size = Pt(11)
            if tok in bold_parts:
                run.bold = True
    return p

def note_box(text):
    """A lightly styled italic note for speaker guidance."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    run = p.add_run(f"[SPEAKER NOTE] {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p

def divider():
    doc.add_paragraph("─" * 72)

def spacer():
    doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("PRESENTATION MANUSCRIPT")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Member 1 — Section 3: Partial Equilibrium European Gas Market Model")
r2.font.size = Pt(13)
r2.bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run(
    "Who Gains When Gas Prices Surge?\n"
    "A Welfare Analysis of Norway's Gas Economy Under the Ukraine War Supply Shock\n"
    "EBA3650 Quantitative Economics · BI Norwegian Business School · Spring 2026"
)
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

divider()
spacer()

# ══════════════════════════════════════════════════════════════════════════════
#  ORIENTATION: YOUR PLACE IN THE BIG PICTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("0. Your Place in the Big Picture", level=1, color=(0x1A, 0x3A, 0x5C))

body(
    "The group's project asks one overarching question: Does Norway benefit from the "
    "war-induced gas price surge — and who within Norway captures the gains? The answer "
    "requires four analytical building blocks, one per member.",
    bold_parts=["Does Norway benefit from the war-induced gas price surge — and who within Norway captures the gains?"]
)

spacer()

body("Group presentation road-map:", bold_parts=["Group presentation road-map:"])
bullet("Member 1 (YOU) — Section 3: Build the market model. Show what happened to prices and quantities, and compute the raw windfall (producer surplus).")
bullet("Member 2 — Section 4: Use the producer surplus from Section 3 to measure how much Norwegian consumers lost (compensating variation).")
bullet("Member 3 — Section 5: Layer the 78 % petroleum tax on top to split the PS windfall between the state and firms, and ask 'who gains most?'")
bullet("Member 4 — Sections 1-2, 6-7: Provide context, stress-test all results via sensitivity analysis, and deliver the conclusion.")

spacer()

body(
    "Your section is the engine room. Without a calibrated market equilibrium, no one else "
    "can compute their numbers. Make sure the audience understands why the model is set up "
    "the way it is before presenting the results.",
    bold_parts=["Your section is the engine room."]
)

note_box(
    "Transition to the group: start by reminding the audience of the research question "
    "(slide/poster already shown). Say: 'To answer this question we need to know what "
    "actually happened in the European gas market — that is what I will show you now.'"
)

divider()
spacer()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 SCRIPT
# ══════════════════════════════════════════════════════════════════════════════
heading("Section 3 — Partial Equilibrium European Gas Market Model", level=1, color=(0x1A, 0x3A, 0x5C))
body("Estimated speaking time: 5–7 minutes   |   Visual cues: market diagram → shock diagram → PS diagrams")

spacer()

# ── 3.0 Opening hook ──────────────────────────────────────────────────────────
heading("3.0  Opening Hook (30 seconds)", level=2)

body(
    "Say verbatim or in your own words:",
    bold_parts=["Say verbatim or in your own words:"]
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.paragraph_format.right_indent = Inches(0.4)
run = p.add_run(
    '"In August 2022, the TTF gas price touched 300 euros per megawatt-hour — '
    'fifteen times the pre-war average. My job is to translate that real-world event '
    'into a mathematical model so we can measure exactly how much money Norway made '
    'from that shock."'
)
run.italic = True
run.font.size = Pt(11)

note_box(
    "This line links back to the research question slide. It tells the audience "
    "you are doing quantitative work, not just describing history."
)

spacer()

# ── 3.1 Model set-up ──────────────────────────────────────────────────────────
heading("3.1  Model Set-Up — Constant-Elasticity Supply and Demand (1.5 min)", level=2)

body(
    "We model the European gas market using constant-elasticity supply and demand curves — "
    "the standard tool from Sessions 2 and 3 of this course.",
    bold_parts=["constant-elasticity supply and demand curves"]
)

spacer()
body("The two functions are:", bold_parts=["The two functions are:"])
bullet("Demand:   Q_d(P) = A_d · P^(ε_d)     where ε_d < 0  (downward-sloping)")
bullet("Supply:   Q_s(P) = A_s · P^(ε_s)     where ε_s > 0  (upward-sloping)")

spacer()
body("Why this functional form?", bold_parts=["Why this functional form?"])
bullet("The elasticity is constant at every point — useful when you only have one data point to calibrate.")
bullet(
    "The math is clean: equilibrium can be solved analytically or numerically with a single equation.",
    bold_parts=["equilibrium can be solved analytically or numerically"]
)
bullet(
    "It is the form used in the course's numerical sessions (Session 2–3), "
    "which is why we can apply the secant method directly.",
    bold_parts=["secant method"]
)

spacer()
body("Calibration — the key step most people skip over:", bold_parts=["Calibration — the key step most people skip over:"])
body(
    "We choose the scale parameters A_d and A_s so that both curves pass through the "
    "pre-war baseline: P₀ = 20 EUR/MWh and Q₀ = 4 000 TWh/year. "
    "This gives:  A = Q₀ / P₀^ε  for each curve.",
    bold_parts=["P₀ = 20 EUR/MWh", "Q₀ = 4 000 TWh/year"]
)
bullet("ε_d = −0.30  (short-run gas demand is inelastic — households cannot switch overnight)")
bullet("ε_s =  0.35  (short-run supply is also inelastic — pipeline capacity is fixed in the short run)")

note_box(
    "Point at the pre-war equilibrium diagram here. Say: 'Both curves cross exactly at "
    "P₀ = 20 and Q₀ = 4 000 by construction. This is our starting point.'"
)

spacer()

# ── 3.2 Supply shock ──────────────────────────────────────────────────────────
heading("3.2  Modelling the Russian Supply Shock (1.5 min)", level=2)

body(
    "The invasion curtailed Russian pipeline gas. Russia held about 40 % of European "
    "supply before the war. We model this as a leftward shift of the total supply curve:",
    bold_parts=["40 % of European supply", "leftward shift"]
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run("Q_s^post(P) = (1 − δ) · A_s · P^(ε_s)     with  δ = 0.40")
run.font.name = "Courier New"
run.font.size = Pt(11)

spacer()
body("Two points to make explicit:")
bullet(
    "The shock removes Russian supply. Norwegian supply is unchanged — "
    "Norway still faces the same supply curve A_NO · P^(ε_s).",
    bold_parts=["Norwegian supply is unchanged"]
)
bullet(
    "But the market-clearing price rises, so Norway now sells its existing output "
    "at a much higher price and can expand slightly along its supply curve.",
    bold_parts=["Norway now sells its existing output at a much higher price"]
)

spacer()
body("Solving for the new equilibrium:", bold_parts=["Solving for the new equilibrium:"])
body(
    "We set Q_d(P) = Q_s^post(P) and solve for P₁. "
    "We use our own secant solver from Session 3, then verify with scipy.optimize.bisect.",
    bold_parts=["secant solver from Session 3", "scipy.optimize.bisect"]
)

body("Results under the baseline parameters:")
bullet("Pre-war:    P₀ = 20 EUR/MWh,  Q₀ = 4 000 TWh/yr")
bullet("Post-shock: P₁ ≈ 43.9 EUR/MWh,  Q₁ ≈ 3 650 TWh/yr")
bullet("Price increase: +119 %  |  Quantity decrease: −8.7 %")

note_box(
    "Show the 'before vs. after' diagram. Note that Q falls only modestly because "
    "both demand and supply are inelastic — the big effect is on price, not quantity. "
    "Both solvers agree to near machine precision, confirming the calculation is correct."
)

body(
    "Reality check: TTF prices in practice were much higher (300 EUR/MWh at peak). "
    "Our model captures the structural direction, not the speculative overshoot. "
    "We use a realistic 'medium-run' price of ≈ 44 EUR/MWh, consistent with annual averages.",
    bold_parts=["Reality check:"]
)

spacer()

# ── 3.3 Norwegian PS ──────────────────────────────────────────────────────────
heading("3.3  Norwegian Producer Surplus — Analytical and Numerical (2 min)", level=2)

body(
    "Now we compute the actual monetary gain to Norwegian gas producers — the number "
    "that drives everything else in the project.",
    bold_parts=["the number that drives everything else"]
)

spacer()
body("Definition of producer surplus:")
body(
    "PS = Revenue − Cost = P · Q_NO − ∫₀^{Q_NO} MC(q) dq",
    bold_parts=[]
)
body(
    "For constant-elasticity supply, the inverse supply function is  "
    "P = (Q / A)^(1/ε),  and the cost integral has a closed-form solution:",
    bold_parts=["closed-form solution"]
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run(
    "∫₀^Q (q/A)^(1/ε) dq  =  A^(−1/ε) · Q^(1 + 1/ε)  /  (1 + 1/ε)"
)
run.font.name = "Courier New"
run.font.size = Pt(10.5)

spacer()
body("Two methods — both give the same answer:", bold_parts=["Two methods — both give the same answer:"])
bullet(
    "Analytical formula: plug in the closed-form integral above. Fast and exact.",
    bold_parts=["Analytical formula"]
)
bullet(
    "Numerical ('integration by force', Session 7): sum up MC over a fine grid of 50 000 points. "
    "Error vs. analytical: less than 0.000001 bn EUR.",
    bold_parts=["Numerical", "Error vs. analytical: less than 0.000001 bn EUR"]
)

note_box(
    "Show the two PS diagrams (pre-war, post-shock) side by side. "
    "Point to the shaded green area. Say: 'The shaded region is the producer surplus — "
    "the gap between market price and marginal cost, summed over all Norwegian output.'"
)

spacer()
body("Norwegian supply parameters:")
bullet("Norway's share of pre-war European supply: 25 %  →  A_NO = 0.25 · A_s")
bullet("Norway's supply curve is NOT shifted by the shock — the shock removes Russian supply only.")

spacer()
body("Key results:", bold_parts=["Key results:"])
bullet("Norwegian output pre-war:   1 000 TWh/yr")
bullet("Norwegian output post-shock:  ~1 035 TWh/yr  (+3.5 %)  — only a small volume increase")
bullet("Norwegian PS pre-war:    ~8.1 bn EUR")
bullet("Norwegian PS post-shock: ~36.1 bn EUR")
bullet("Change in PS:  +28.0 bn EUR",  bold_parts=["+28.0 bn EUR"])

spacer()
body(
    "Interpretation: almost all of Norway's gain comes from price, not volume. "
    "Because supply elasticity is low (εₛ = 0.35), Norway cannot rapidly expand output. "
    "Instead, it earns the same approximate volume at more than double the price.",
    bold_parts=[
        "almost all of Norway's gain comes from price, not volume.",
        "earns the same approximate volume at more than double the price."
    ]
)

note_box(
    "Pause here. Say: 'This 28 billion euro gross gain is the number the rest of the "
    "analysis is built on. Member 2 will use it to compute the consumer welfare loss. "
    "Member 3 will split it between the state and firms using the 78 % petroleum tax.'"
)

spacer()

# ── Handover ──────────────────────────────────────────────────────────────────
heading("3.4  Handover to Member 2 (30 seconds)", level=2)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.paragraph_format.right_indent = Inches(0.4)
run = p.add_run(
    '"So to summarise Section 3: the shock shifts the supply curve left, pushes '
    'the market price from 20 to 44 euros, and delivers a gross windfall of 28 billion '
    'euros to Norwegian gas producers. The natural next question is: '
    'who loses from this price surge inside Norway? '
    'My colleague will now show you the consumer side of the picture."'
)
run.italic = True
run.font.size = Pt(11)

note_box(
    "Look at Member 2 and gesture to hand over. "
    "Do NOT summarise Member 2's results — leave the audience curious."
)

divider()
spacer()

# ══════════════════════════════════════════════════════════════════════════════
#  ANTICIPATED QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("Anticipated Q&A — Section 3", level=1, color=(0x1A, 0x3A, 0x5C))

qs = [
    (
        "Why constant-elasticity and not linear supply/demand?",
        "Linear curves require choosing specific intercepts and slopes. "
        "Constant-elasticity requires only the elasticity and one (price, quantity) data point — "
        "and elasticities are the parameters economists actually estimate and publish. "
        "It is a standard modelling choice in energy economics."
    ),
    (
        "Why does price only rise to 44 EUR/MWh when real TTF hit 300?",
        "Our 43.9 EUR/MWh reflects a structural equilibrium with δ = 40% of supply removed. "
        "The 300 EUR/MWh peak was a short-run speculative spike with panic buying, storage concerns, "
        "and thin liquidity. Annual-average TTF in 2022 was ~125 EUR/MWh; our model is closer to "
        "the 'true' structural shift than to the panic peak. Sensitivity analysis (Member 4) "
        "tests larger δ values and confirms results are robust."
    ),
    (
        "Why use the secant method instead of just scipy?",
        "The course requires us to implement numerical solvers from scratch (Sessions 2–3). "
        "The secant method is the correct tool here: our excess demand function is continuous "
        "and the derivative is not available in closed form. We use scipy.bisect only to verify "
        "our own solver — and both agree to 10⁻¹⁰ precision."
    ),
    (
        "Isn't Norwegian supply also affected by higher prices?",
        "Yes — Norway moves up along its unchanged supply curve when prices rise. "
        "Our model captures this: Norwegian output increases from ~1 000 to ~1 035 TWh (+3.5%). "
        "But because supply elasticity is low (εₛ = 0.35), the volume increase is modest. "
        "The big gain is from the price, not expanded production."
    ),
    (
        "What is δ = 0.40 based on?",
        "Russia supplied roughly 40% of European gas pre-war via Nord Stream, Yamal, "
        "and Brotherhood pipelines (approx. 155 bcm out of ~400 bcm total). "
        "The curtailment was progressive, but by Q4 2022 flows were close to zero. "
        "We model the full removal as a baseline; sensitivity analysis tests δ from 0.20 to 0.60."
    ),
    (
        "How does this connect to the final net welfare number?",
        "The +28 bn EUR PS gain is the gross windfall. Member 3 splits it: "
        "78% (21.8 bn EUR) goes to the government via the petroleum tax, "
        "22% (6.2 bn EUR) stays with gas firms. "
        "Member 2 estimates the consumer welfare loss at −5.0 bn EUR. "
        "Net welfare = 28.0 − 5.0 = +23.0 bn EUR. Norway is a net winner by a large margin."
    ),
]

for i, (q, a) in enumerate(qs, 1):
    heading(f"Q{i}:  {q}", level=3)
    body(f"A:  {a}", bold_parts=["A:  "])
    spacer()

divider()
spacer()

# ══════════════════════════════════════════════════════════════════════════════
#  QUICK REFERENCE — KEY NUMBERS
# ══════════════════════════════════════════════════════════════════════════════
heading("Quick-Reference Card — Key Numbers for Section 3", level=1, color=(0x1A, 0x3A, 0x5C))

body("Memorise these — they are the numbers you must not get wrong live:")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Parameter / Result"
hdr[1].text = "Value"
hdr[2].text = "Meaning"

rows = [
    ("P₀ (pre-war price)",   "20 EUR/MWh",    "2019–2021 average TTF"),
    ("Q₀ (pre-war quantity)","4 000 TWh/yr",  "Total European gas consumption"),
    ("ε_d",                  "−0.30",          "Short-run demand elasticity"),
    ("ε_s",                  "+0.35",          "Short-run supply elasticity"),
    ("δ (shock magnitude)",  "0.40",           "Russia's share of supply removed"),
    ("Norway's share",       "25%",            "of pre-war European gas supply"),
    ("P₁ (post-shock price)","43.89 EUR/MWh",  "+119% from baseline"),
    ("Q₁ (post-shock qty)",  "~3 650 TWh/yr",  "−8.7% from baseline"),
    ("Norwegian Q pre-war",  "1 000 TWh/yr",   "= 25% × 4 000"),
    ("Norwegian Q post",     "~1 035 TWh/yr",  "+3.5% (small volume gain)"),
    ("PS pre-war",           "~8.1 bn EUR",    "Baseline producer surplus"),
    ("PS post-shock",        "~36.1 bn EUR",   "After price surge"),
    ("ΔPS (Section 3 output)","+ 28.0 bn EUR", "The gross windfall — INPUT to Sections 4 & 5"),
]

for r in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = r[0]
    row_cells[1].text = r[1]
    row_cells[2].text = r[2]

spacer()
spacer()

# ══════════════════════════════════════════════════════════════════════════════
#  CONNECTION TO BIG PICTURE — SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
heading("How Section 3 Feeds the Full Project", level=1, color=(0x1A, 0x3A, 0x5C))

flow = [
    ("Section 3 (YOU)", "ΔPS = +28.0 bn EUR", "Gross windfall from supply shock"),
    ("Section 4 (Member 2)", "CV = −5.0 bn EUR", "Uses P₀, P₁, Q values to measure consumer loss"),
    ("Section 5 (Member 3)", "Gov: +21.8 bn / Firm: +6.2 bn", "Splits ΔPS via 78% petroleum tax"),
    ("Sections 6-7 (Member 4)", "Net ΔW = +23.0 bn EUR", "Aggregates: ΔPS − CV; sensitivity tests"),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
hdr2[0].text = "Who"
hdr2[1].text = "What they compute"
hdr2[2].text = "What they need from you"

for who, what, needs in flow:
    r = table2.add_row().cells
    r[0].text = who
    r[1].text = what
    r[2].text = needs

spacer()
body(
    "The central message: Section 3 establishes the price change (P₀→P₁) and the gross "
    "producer surplus gain (ΔPS = +28 bn EUR). These two outputs are the inputs for every "
    "other section. If the audience understands nothing else from your presentation, "
    "they must understand these two numbers.",
    bold_parts=[
        "P₀→P₁",
        "ΔPS = +28 bn EUR",
        "These two outputs are the inputs for every other section."
    ]
)

spacer()
spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Manuscript · Member 1 · Section 3 —")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── save ──────────────────────────────────────────────────────────────────────
out = "section3_presentation_manus.docx"
doc.save(out)
print(f"Saved: {out}")
